"""Xuất báo cáo học phần ra .docx khổ A4 chuẩn nộp.

Nguyên tắc thiết kế: **mọi thứ đều đọc từ file Markdown**, không có danh sách
cứng nào trong mã. Chú thích bảng và hình được lấy từ chính nhãn ``*Bảng N — …*``
và ``*Hình N — …*`` trong báo cáo, còn sơ đồ lấy từ PNG do ``render_so_do.py``
sinh ra trực tiếp từ khối ```mermaid hoặc ```plantuml. Nhờ vậy bản xuất không
thể lệch khỏi bản Markdown — sửa báo cáo rồi chạy lại là khớp.

Cách dùng::

    python docs/bao-cao/render_so_do.py      # sinh lại 4 sơ đồ
    python docs/bao-cao/xuat_bao_cao_docx.py # xuất .docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

# ----------------------------------------------------------------------------
# Cấu hình
# ----------------------------------------------------------------------------

HERE = Path(__file__).resolve().parent
ROOT = HERE.parents[1]
import sys
# Tên báo cáo BẮT BUỘC truyền vào. Trước đây chỗ này mặc định về một tệp cụ thể; khi tệp đó bị xoá
# thì script chết bằng traceback của thư viện, ở tận chỗ đọc nội dung — xa nguyên nhân thật.
if len(sys.argv) <= 1:
    sys.exit("Cần tên tệp báo cáo, ví dụ:\n"
             "    python docs/bao-cao/xuat_bao_cao_docx.py BAO_CAO_<HỌC_PHẦN>.md\n"
             "Chưa có báo cáo nào thì chép `KHUON_BAO_CAO.md` ra làm bản đầu.")
NGUON = HERE / sys.argv[1]
if not NGUON.is_file():
    sys.exit(f"Không thấy {NGUON.relative_to(ROOT).as_posix()}")
THU_MUC_RA = HERE / "output"
DICH = THU_MUC_RA / (NGUON.stem + ".docx")
SO_DO = THU_MUC_RA / ("_diagrams_" + NGUON.stem[:18])

REPO = "https://github.com/Anpham120/restaurant-qr-ordering-mobile"
NHANH = "develop"

FONT = "Times New Roman"
CO_CHU = 13
GIAN_DONG = 1.5
FONT_MA = "Consolas"

LE_TRAI, LE_PHAI, LE_TREN, LE_DUOI = 3.0, 2.0, 2.0, 2.0
RONG_NOI_DUNG_CM = 21.0 - LE_TRAI - LE_PHAI          # 16 cm

# --- Công tắc bố cục ------------------------------------------------------
# Để False khi ghép nội dung vào một mẫu bìa đã có sẵn đầu trang / số trang,
# tránh việc hai nguồn cùng sinh ra một thứ.
CO_DAU_TRANG = False      # dòng chữ chạy ở đầu mỗi trang
CO_SO_TRANG = False       # số trang ở chân trang
CO_TRANG_BIA = True       # trang bìa do script tự dựng

# Lấy nguyên trang bìa, đầu trang và chân trang của một tệp DOCX đã có sẵn.
# Khi bật, script mở tệp mẫu, giữ lại đúng khối trang bìa rồi xóa phần thân,
# sau đó dựng nội dung mới vào. Nhờ vậy logo ở bìa và ở đầu trang, số trang ở
# chân trang đều giữ nguyên mà không phải dựng lại bằng mã.
# Đặt None để quay về trang bìa do script tự dựng.
MAU_BIA = HERE / "output" / "BAO_CAO_PHAN_MEM_HOAN_CHINH.docx"
# Đoạn đầu tiên KHÔNG thuộc trang bìa trong tệp mẫu. Mọi thứ từ đây trở đi
# bị xóa để nhường chỗ cho nội dung mới.
MAU_BIA_KET_THUC = "MỤC LỤC"

# Chỉ đề mục được in đậm. Mọi chỗ khác — thân bài, ô bảng, dòng tiêu đề bảng,
# chú thích, trang bìa — đều về chữ thường, kể cả khi Markdown có đánh dấu **.
# Dấu ** trong Markdown vẫn được giữ nguyên để đảo lại được bằng cách đặt False.
CHI_DE_MUC_DAM = True
# Các tiêu đề không nằm trong cây đề mục nhưng đóng vai trò như đề mục, nên
# vẫn giữ in đậm khi CHI_DE_MUC_DAM bật.
DE_MUC_NGOAI_CAY = ("MỤC LỤC", "DANH MỤC BẢNG", "DANH MỤC HÌNH",
                    "DANH MỤC HÌNH VẼ", "DANH MỤC TỪ VIẾT TẮT")

# Chiều cao TỐI THIỂU của hàng trong bảng. Hàng nhiều chữ vẫn tự giãn thêm.
CAO_HANG_TIEU_DE_CM = 0.9
CAO_HANG_CM = 0.75

MUC = RGBColor(0x44, 0x44, 0x44)
DEN = RGBColor(0x00, 0x00, 0x00)
# Mau chu cua moi de muc. Bao cao in den trang nen dat ve den; doi bien
# nay neu muon de muc co mau.
MAU_DE_MUC = DEN
XANH_LIEN_KET = RGBColor(0x1F, 0x4E, 0x79)   # chi dung cho lien ket
VIEN = "BFBFBF"
NEN_TIEU_DE_BANG = "EDEDED"
NEN_MA = "F4F4F4"


# ----------------------------------------------------------------------------
# Tiện ích XML mức thấp
# ----------------------------------------------------------------------------

def _el(tag: str, **attrs) -> OxmlElement:
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), str(v))
    return e


def to_mau_nen(o, mau: str) -> None:
    """Tô nền cho ô bảng hoặc đoạn văn."""
    pr = o._tc.get_or_add_tcPr() if hasattr(o, "_tc") else o._p.get_or_add_pPr()
    pr.append(_el("w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": mau}))


def dat_vien_bang(bang) -> None:
    pr = bang._tbl.tblPr
    borders = _el("w:tblBorders")
    for canh in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_el(f"w:{canh}", **{"w:val": "single", "w:sz": "4",
                                           "w:space": "0", "w:color": VIEN}))
    pr.append(borders)


def lap_lai_dong_tieu_de(dong) -> None:
    dong._tr.get_or_add_trPr().append(_el("w:tblHeader", **{"w:val": "true"}))


def khong_tach_dong(dong) -> None:
    dong._tr.get_or_add_trPr().append(_el("w:cantSplit"))


def cao_toi_thieu(dong, cm: float) -> None:
    """Đặt chiều cao *tối thiểu* cho hàng, không phải chiều cao cố định.

    Dùng quy tắc ``atLeast`` để mọi hàng cao bằng nhau khi nội dung ngắn, còn
    hàng nào nhiều chữ thì tự giãn ra. Nếu đặt ``exact`` thì Word sẽ cắt cụt
    phần chữ vượt quá, đúng thứ cần tránh.
    """
    tr = dong._tr.get_or_add_trPr()
    tr.append(_el("w:trHeight", **{"w:val": str(int(cm * 567)),   # cm -> twip
                                   "w:hRule": "atLeast"}))


def them_truong(doan, ma: str) -> None:
    """Chèn một Word field (dùng cho mục lục và số trang)."""
    r1, r2, r3 = doan.add_run(), doan.add_run(), doan.add_run()
    r1._r.append(_el("w:fldChar", **{"w:fldCharType": "begin"}))
    t = _el("w:instrText", **{"xml:space": "preserve"})
    t.text = ma
    r2._r.append(t)
    r3._r.append(_el("w:fldChar", **{"w:fldCharType": "end"}))


def them_lien_ket(doan, chu: str, url: str):
    rid = doan.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    h = _el("w:hyperlink", **{"r:id": rid})
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = _el("w:color", **{"w:val": "1F4E79"})
    u = _el("w:u", **{"w:val": "single"})
    rpr.append(c)
    rpr.append(u)
    r.append(rpr)
    t = _el("w:t", **{"xml:space": "preserve"})
    t.text = chu
    r.append(t)
    h.append(r)
    doan._p.append(h)


# ----------------------------------------------------------------------------
# Kiểu chữ
# ----------------------------------------------------------------------------

def dung_kieu(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = FONT
    n.font.size = Pt(CO_CHU)
    n._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = n.paragraph_format
    pf.line_spacing = GIAN_DONG
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.widow_control = True          # không để một dòng lẻ trôi sang trang sau

    # Quy ước trình bày Việt Nam khi chữ thường 13pt: cỡ chữ đề mục gần như
    # không đổi, phân cấp bằng IN HOA / đậm / nghiêng và bằng số thứ tự.
    #   Heading 1  14pt  IN HOA, đậm, căn giữa   (tên chương)
    #   Heading 2  13pt  đậm                     (1.1)
    #   Heading 3  13pt  đậm, thụt lề 0,5 cm     (1.1.1)
    #   Heading 4  13pt  đậm, thụt lề 0,5 cm     (mục phụ, không đánh số)
    # Mọi đề mục chỉ dùng in đậm — không nghiêng. Phân cấp thể hiện bằng số
    # thứ tự và mức thụt lề.
    for ten, co, dam, nghieng, hoa, can, thut, truoc, sau in (
            ("Heading 1", 14, True, False, True, WD_ALIGN_PARAGRAPH.CENTER, 0.0, 20, 12),
            ("Heading 2", 13, True, False, False, WD_ALIGN_PARAGRAPH.LEFT, 0.0, 14, 8),
            ("Heading 3", 13, True, False, False, WD_ALIGN_PARAGRAPH.LEFT, 0.5, 12, 6),
            ("Heading 4", 13, True, False, False, WD_ALIGN_PARAGRAPH.LEFT, 0.5, 10, 6)):
        s = doc.styles[ten]
        s.font.name = FONT
        s.font.size = Pt(co)
        s.font.bold = dam
        s.font.italic = nghieng
        s.font.color.rgb = MAU_DE_MUC
        if hoa:
            s.element.get_or_add_rPr().append(_el("w:caps", **{"w:val": "true"}))
        s.paragraph_format.space_before = Pt(truoc)
        s.paragraph_format.space_after = Pt(sau)
        s.paragraph_format.line_spacing = 1.3
        s.paragraph_format.alignment = can
        s.paragraph_format.left_indent = Cm(thut)
        s.paragraph_format.keep_with_next = True


def dat_trang(doc: Document) -> None:
    for s in doc.sections:
        s.page_width, s.page_height = Cm(21), Cm(29.7)
        s.left_margin, s.right_margin = Cm(LE_TRAI), Cm(LE_PHAI)
        s.top_margin, s.bottom_margin = Cm(LE_TREN), Cm(LE_DUOI)


def dat_dau_chan(doc: Document, tieu_de: str) -> None:
    """Đặt đầu trang và số trang theo công tắc ở đầu tệp.

    Mặc định cả hai đều tắt, vì nội dung được ghép vào mẫu bìa đã có sẵn phần
    này. Khi bật, trang bìa vẫn được để trống và số trang đếm lại từ 1 kể từ
    trang mục lục.
    """
    for k, s in enumerate(doc.sections):
        s.header.is_linked_to_previous = False
        s.footer.is_linked_to_previous = False
        s.header.paragraphs[0].text = ""
        s.footer.paragraphs[0].text = ""

        if k == 0:                      # section chứa trang bìa: luôn để trống
            continue

        if CO_DAU_TRANG:
            h = s.header.paragraphs[0]
            h.text = tieu_de
            h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in h.runs:
                r.font.size = Pt(10)
                r.font.name = FONT
                r.font.color.rgb = MUC
                r.font.italic = True

        if CO_SO_TRANG:
            f = s.footer.paragraphs[0]
            f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            them_truong(f, "PAGE")
            for r in f.runs:
                r.font.size = Pt(11)
                r.font.name = FONT
            if k == 1:                  # bắt đầu đánh số lại từ 1
                s._sectPr.append(_el("w:pgNumType", **{"w:start": "1"}))


# ----------------------------------------------------------------------------
# Xử lý văn bản nội tuyến
# ----------------------------------------------------------------------------

MA_INLINE = re.compile(r"`([^`]+)`")
DAM = re.compile(r"\*\*([^*]+)\*\*")
NGHIENG = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
LIEN_KET = re.compile(r"\[([^\]]*)\]\(([^)]+)\)")


def doi_duong_dan(dich: str) -> str:
    """Đổi đường dẫn tương đối trong repo thành URL GitHub."""
    if dich.startswith(("http://", "https://", "mailto:")):
        return dich
    if dich.startswith("#"):
        return ""
    duong = (HERE / dich.split("#")[0]).resolve()
    try:
        rel = duong.relative_to(ROOT).as_posix()
    except ValueError:
        return ""
    loai = "tree" if duong.is_dir() else "blob"
    return f"{REPO}/{loai}/{NHANH}/{rel}"


def viet_inline(doan, text: str, dam=False, nghieng=False, co=None) -> None:
    """Ghi một chuỗi Markdown nội tuyến vào đoạn văn."""
    text = re.sub(r"<br\s*/?>", "\n", text)
    text = re.sub(r"</?(sub|strong|em|b|i|code)>", "", text)

    vi_tri = 0
    # Nhãn của liên kết được phép chứa một cấp ngoặc vuông lồng nhau, để bắt
    # đúng dạng dấu trích dẫn [[1]](url) và [[1, Ch.2]](url) trong báo cáo.
    mau = re.compile(
        r"\[((?:[^\[\]]|\[[^\[\]]*\])*)\]\(([^)]+)\)"   # 1,2 liên kết
        r"|\*\*([^*]+)\*\*"              # 3 đậm
        r"|(?<!\*)\*([^*]+)\*(?!\*)"     # 4 nghiêng
        r"|`([^`]+)`"                    # 5 mã
    )
    for m in mau.finditer(text):
        if m.start() > vi_tri:
            _run(doan, text[vi_tri:m.start()], dam, nghieng, co)
        if m.group(1) is not None:
            url = doi_duong_dan(m.group(2))
            nhan = m.group(1) or m.group(2)
            if url:
                them_lien_ket(doan, nhan, url)
            else:
                _run(doan, nhan, dam, nghieng, co)
        elif m.group(3) is not None:
            # Nội dung in đậm vẫn có thể chứa liên kết -> xử lý đệ quy.
            if "](" in m.group(3):
                viet_inline(doan, m.group(3), True, nghieng, co)
            else:
                _run(doan, m.group(3), True, nghieng, co)
        elif m.group(4) is not None:
            if "](" in m.group(4):
                viet_inline(doan, m.group(4), dam, True, co)
            else:
                _run(doan, m.group(4), dam, True, co)
        else:
            r = _run(doan, m.group(5), dam, nghieng, co)
            r.font.name = FONT_MA
            r.font.size = Pt((co or CO_CHU) - 2)
        vi_tri = m.end()
    if vi_tri < len(text):
        _run(doan, text[vi_tri:], dam, nghieng, co)


def _run(doan, chu: str, dam: bool, nghieng: bool, co):
    if not chu:
        return doan.add_run("")
    dong = chu.split("\n")
    r = None
    for k, phan in enumerate(dong):
        if k:
            r = doan.add_run()
            r.add_break(WD_BREAK.LINE)
        r = doan.add_run(phan)
        r.bold = dam
        r.italic = nghieng
        r.font.name = FONT
        if co:
            r.font.size = Pt(co)
    return r


# ----------------------------------------------------------------------------
# Các khối
# ----------------------------------------------------------------------------

def them_anh(doc: Document, duong: Path, rong_toi_da_cm=None, cao_toi_da_cm=22.0):
    if not duong.exists():
        raise FileNotFoundError(duong)
    rong_toi_da_cm = rong_toi_da_cm or RONG_NOI_DUNG_CM
    with Image.open(duong) as im:
        px_w, px_h = im.size
    rong = rong_toi_da_cm
    cao = rong * px_h / px_w
    if cao > cao_toi_da_cm:
        cao = cao_toi_da_cm
        rong = cao * px_w / px_h
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(duong), width=Cm(rong))
    return p


def doi_huong_trang(doc: Document, ngang: bool):
    """Mở một section mới với hướng trang mong muốn.

    Dùng cho sơ đồ quá rộng: đặt nằm ngang thì chiều rộng dùng được tăng từ
    16 cm lên 25,7 cm, tức lớn hơn 1,6 lần.
    """
    from docx.enum.section import WD_ORIENT

    s = doc.add_section(WD_SECTION.NEW_PAGE)
    if ngang:
        s.orientation = WD_ORIENT.LANDSCAPE
        s.page_width, s.page_height = Cm(29.7), Cm(21.0)
        s.left_margin = s.right_margin = Cm(2.0)
        s.top_margin = s.bottom_margin = Cm(2.0)
    else:
        s.orientation = WD_ORIENT.PORTRAIT
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        s.left_margin, s.right_margin = Cm(LE_TRAI), Cm(LE_PHAI)
        s.top_margin, s.bottom_margin = Cm(LE_TREN), Cm(LE_DUOI)
    return s


def them_so_do(doc: Document, duong: Path, nguong_ngang=2.0):
    """Chèn sơ đồ, tự chuyển sang trang ngang nếu nó quá rộng."""
    with Image.open(duong) as im:
        w, h = im.size
    ti_le = w / h
    if ti_le >= nguong_ngang:
        doi_huong_trang(doc, ngang=True)
        p = them_anh(doc, duong, rong_toi_da_cm=25.7, cao_toi_da_cm=16.0)
        return p, True
    return them_anh(doc, duong, cao_toi_da_cm=19.5), False


def them_chu_thich(doc: Document, chu: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_with_next = False
    viet_inline(p, chu, nghieng=True, co=11.5)
    for r in p.runs:
        r.italic = True
        r.font.color.rgb = DEN
    return p


def them_khoi_ma(doc: Document, ma: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    to_mau_nen(p, NEN_MA)
    for k, dong in enumerate(ma.rstrip().split("\n")):
        r = p.add_run()
        if k:
            r.add_break(WD_BREAK.LINE)
        r = p.add_run(dong)
        r.font.name = FONT_MA
        r.font.size = Pt(9.5)
    return p


def them_trich_dan(doc: Document, dong_van: list[str]):
    for dong in dong_van:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3
        to_mau_nen(p, "F7F9FC")
        viet_inline(p, dong, co=12)
        for r in p.runs:
            if r.font.size is None:
                r.font.size = Pt(12)


def tach_o(dong: str) -> list[str]:
    dong = dong.strip()
    if dong.startswith("|"):
        dong = dong[1:]
    if dong.endswith("|"):
        dong = dong[:-1]
    return [c.strip() for c in re.split(r"(?<!\\)\|", dong)]


def _chu_thuan(x: str) -> str:
    """Bỏ cú pháp Markdown để đo đúng số ký tự sẽ hiển thị."""
    x = re.sub(r"\[([^\]]*)\]\([^)]+\)", r"\1", x)   # liên kết -> nhãn
    x = re.sub(r"<br\s*/?>", " ", x)
    x = re.sub(r"[*`<>]", "", x)
    return x.strip()


def rong_cot_theo_noi_dung(tieu_de: list[str], than: list[list[str]],
                           co_chu: float) -> list[float]:
    """Chia chiều rộng cột sao cho không cột nào hẹp hơn từ dài nhất của nó.

    Cách chia gồm hai bước. Trước hết mỗi cột được cấp một chiều rộng **tối
    thiểu** đủ chứa từ dài nhất xuất hiện trong cột đó, kể cả ở dòng tiêu đề —
    nếu thiếu bước này Word sẽ ngắt giữa từ và hàng tiêu đề bị đội cao. Phần
    chỗ còn dư mới đem chia tiếp theo lượng chữ trung bình, nên cột nhiều nội
    dung được rộng hơn mà cột ngắn vẫn đủ chỗ.
    """
    n = len(tieu_de)
    rong_ky_tu = 0.0176 * co_chu        # bề rộng trung bình một ký tự, cm
    dem = 0.36                          # đệm trái + phải của ô
    # Trần nới rộng hơn cho bảng ít cột, vì ở đó một định danh mã dài vẫn còn
    # chỗ mà không bóp nghẹt các cột còn lại.
    tran_toi_thieu = RONG_NOI_DUNG_CM * (0.46 if n <= 3 else 0.34)

    toi_thieu, mong_muon = [], []
    for i in range(n):
        nhan = _chu_thuan(tieu_de[i])
        o = [nhan] + [_chu_thuan(h[i]) for h in than if i < len(h)]

        # (a) không ngắt giữa từ dài nhất, kể cả từ trong tiêu đề
        tu = [w for x in o for w in re.split(r"[\s/,;()]+", x) if w]
        dai_nhat = max((len(w) for w in tu), default=1)
        rong_tu = dai_nhat * rong_ky_tu + dem

        # (b) tiêu đề cột gọn trong tối đa hai dòng, không vỡ vụn thành 3–4 dòng
        rong_nhan = (len(nhan) / 2 + 1.5) * rong_ky_tu + dem

        toi_thieu.append(min(max(rong_tu, rong_nhan), tran_toi_thieu))

        tb = sum(len(x) for x in o) / len(o)
        mong_muon.append(tb * rong_ky_tu * 0.62 + dem)

    tong_min = sum(toi_thieu)
    if tong_min >= RONG_NOI_DUNG_CM:    # bảng quá chật: co đều theo tỉ lệ
        return [x * RONG_NOI_DUNG_CM / tong_min for x in toi_thieu]

    du = RONG_NOI_DUNG_CM - tong_min
    them = [max(mong_muon[i] - toi_thieu[i], 0) for i in range(n)]
    tong_them = sum(them)
    if tong_them <= 0:
        return [x + du / n for x in toi_thieu]
    return [toi_thieu[i] + du * them[i] / tong_them for i in range(n)]


def can_le_cot(than: list[list[str]], i: int) -> int:
    """Mọi ô trong bảng đều căn giữa cho cân đối."""
    return WD_ALIGN_PARAGRAPH.CENTER


def them_bang(doc: Document, dong_bang: list[str]):
    tieu_de = tach_o(dong_bang[0])
    than = [tach_o(d) for d in dong_bang[2:] if d.strip()]
    so_cot = len(tieu_de)
    bang = doc.add_table(rows=1, cols=so_cot)
    bang.autofit = False
    dat_vien_bang(bang)

    co_chu_bang = 11 if so_cot <= 4 else (10 if so_cot == 5 else 9)
    rong = rong_cot_theo_noi_dung(tieu_de, than, co_chu_bang)
    can_le = [can_le_cot(than, i) for i in range(so_cot)]
    for i, c in enumerate(bang.columns):
        c.width = Cm(rong[i])

    for i, chu in enumerate(tieu_de):
        o = bang.rows[0].cells[i]
        o.width = Cm(rong[i])
        o.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        to_mau_nen(o, NEN_TIEU_DE_BANG)
        p = o.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        viet_inline(p, chu, dam=True, co=co_chu_bang)
        for r in p.runs:
            r.bold = True
    lap_lai_dong_tieu_de(bang.rows[0])
    khong_tach_dong(bang.rows[0])
    cao_toi_thieu(bang.rows[0], CAO_HANG_TIEU_DE_CM)

    for hang in than:
        o_moi = bang.add_row()
        khong_tach_dong(o_moi)
        cao_toi_thieu(o_moi, CAO_HANG_CM)
        for i in range(so_cot):
            o = o_moi.cells[i]
            o.width = Cm(rong[i])   # Word đọc bề rộng ở mức ô, không ở mức cột
            o.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = o.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = can_le[i]
            viet_inline(p, hang[i] if i < len(hang) else "", co=co_chu_bang)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return bang


def them_thu_vien_anh(doc: Document, html: str):
    """Bảng ảnh dạng <table> trong Markdown -> bảng Word chứa ảnh."""
    hang_html = re.findall(r"<tr\b.*?</tr>", html, flags=re.S | re.I)
    du_lieu = []
    for h in hang_html:
        o = re.findall(r"<td\b.*?</td>", h, flags=re.S | re.I)
        du_lieu.append(o)
    so_cot = max(len(h) for h in du_lieu)

    bang = doc.add_table(rows=0, cols=so_cot)
    bang.autofit = True
    rong_o = (RONG_NOI_DUNG_CM - 0.6) / so_cot

    for o_hang in du_lieu:
        hang = bang.add_row()
        khong_tach_dong(hang)
        for i, o_html in enumerate(o_hang):
            o = hang.cells[i]
            o.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = o.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)

            m = re.search(r'<img[^>]*src="([^"]+)"', o_html)
            if m:
                duong = (HERE / m.group(1)).resolve()
                if duong.exists():
                    with Image.open(duong) as im:
                        pw, ph = im.size
                    rong = rong_o
                    if rong * ph / pw > 11.0:
                        rong = 11.0 * pw / ph
                    p.add_run().add_picture(str(duong), width=Cm(rong))

            for the, dam, co in (("strong", True, 11.5), ("sub", False, 10)):
                for chu in re.findall(rf"<{the}>(.*?)</{the}>", o_html, flags=re.S | re.I):
                    q = o.add_paragraph()
                    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    q.paragraph_format.space_after = Pt(1)
                    q.paragraph_format.line_spacing = 1.1
                    viet_inline(q, chu.strip(), dam=dam, co=co)
                    for r in q.runs:
                        r.font.size = Pt(co)
                        r.bold = dam
                        if not dam:
                            r.font.color.rgb = MUC
                            r.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return bang


# ----------------------------------------------------------------------------
# Trang bìa
# ----------------------------------------------------------------------------

def trang_bia(doc: Document, meta: dict) -> None:
    def d(chu, co, dam=False, truoc=0, sau=6, mau=None, hoa=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(truoc)
        p.paragraph_format.space_after = Pt(sau)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(chu.upper() if hoa else chu)
        r.bold = dam
        r.font.name = FONT
        r.font.size = Pt(co)
        if mau:
            r.font.color.rgb = mau
        return p

    d("TRƯỜNG ĐẠI HỌC CMC", 14, True, 0, 2)
    d("KHOA CÔNG NGHỆ THÔNG TIN & TRUYỀN THÔNG", 13, True, 0, 24)

    logo = (HERE / "../../frontend/src/mocks/images/logo.png").resolve()
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(20)
        p.add_run().add_picture(str(logo), width=Cm(3.6))

    d(meta.get("loai", "BÁO CÁO"), 20, True, 0, 6, MAU_DE_MUC, True)
    if meta.get("hoc_phan"):
        d(meta["hoc_phan"], 13, False, 0, 26)
    else:
        d("", 13, False, 0, 20)

    d("ĐỀ TÀI", 13, True, 0, 6)
    d(meta["de_tai"], 16, True, 0, 30, MAU_DE_MUC)

    for nhan, gia in (("Giảng viên hướng dẫn:", meta["gv"]),
                      ("Nhóm thực hiện:", meta["nhom"]),
                      ("Thời gian thực hiện:", meta["thoi_gian"])):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{nhan} ")
        r.font.size = Pt(13)
        r.font.name = FONT
        r = p.add_run(gia)
        r.bold = True
        r.font.size = Pt(13)
        r.font.name = FONT

    doc.add_paragraph()
    d(meta["dia_diem"], 13, False, 30, 0)
    # Trang bìa là một section riêng để không mang đầu trang và số trang.
    doi_huong_trang(doc, ngang=False)


# ----------------------------------------------------------------------------
# Mục lục tự động
# ----------------------------------------------------------------------------

# Số đoạn đầu tài liệu thuộc trang bìa lấy theo mẫu. Dùng chỉ số chứ không dùng
# id() của phần tử XML: lxml tạo lại đối tượng bao mỗi lần truy cập nên id()
# không ổn định giữa hai lần duyệt.
_SO_DOAN_BIA = 0


def mo_theo_mau(meta: dict) -> Document:
    """Mở tệp mẫu, giữ trang bìa, xóa phần thân, rồi thay thông tin học phần.

    Đầu trang và chân trang nằm ở phần riêng của section nên tự động đi theo,
    kể cả ảnh và trường số trang. Trang bìa chỉ đổi ba chỗ: tên loại báo cáo,
    tên học phần và tên giảng viên; phần còn lại — logo, tên dự án, danh sách
    sinh viên, ngày tháng — giữ nguyên vì dùng chung cho mọi báo cáo.
    """
    doc = Document(str(MAU_BIA))
    than = doc.element.body

    # Dùng Paragraph.text chứ không phải itertext(): itertext() gom cả phần nội
    # dung dự phòng của mc:AlternateContent nên chuỗi bị nhân đôi.
    moc = next((p._element for p in doc.paragraphs
                if p.text.strip() == MAU_BIA_KET_THUC), None)
    if moc is None:
        raise SystemExit(f"Không tìm thấy '{MAU_BIA_KET_THUC}' trong {MAU_BIA.name}")

    xoa = False
    for el in list(than):
        if el is moc:
            xoa = True
        if xoa and el.tag != qn("w:sectPr"):   # giữ sectPr cuối tài liệu
            than.remove(el)

    # Trang bìa của mẫu để chữ đậm; giữ nguyên cỡ và căn lề, chỉ thay nội dung.
    def thay_ca_doan(p, chu):
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
        p.runs[0].text = chu

    def thay_mot_cum(p, cu, moi):
        """Thay một chuỗi con dù Word đã cắt nó thành nhiều run.

        Word thường tách run theo rsid nên chuỗi cần thay hiếm khi nằm gọn
        trong một run. Ở đây dồn chữ mới vào run đầu tiên chạm tới chuỗi cũ
        rồi cắt phần thừa ở các run sau, nhờ vậy giữ được định dạng gốc.
        """
        vt = p.text.find(cu)
        if vt < 0:
            return False
        het = vt + len(cu)
        moc_dau = 0
        for r in p.runs:
            moc_cuoi = moc_dau + len(r.text)
            if moc_cuoi > vt and moc_dau < het:
                dau = r.text[:max(vt - moc_dau, 0)]
                duoi = r.text[max(het - moc_dau, 0):]
                r.text = dau + (moi if moc_dau <= vt else "") + duoi
            moc_dau = moc_cuoi
        return True

    doi = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("BÁO CÁO") and meta["loai"]:
            thay_ca_doan(p, meta["loai"])
            doi.append("loại báo cáo")
        elif t.startswith("MÔN:") and meta["hoc_phan"]:
            thay_ca_doan(p, "MÔN: " + meta["hoc_phan"].replace("Học phần:", "").strip())
            doi.append("học phần")
        elif "Giảng viên hướng dẫn:" in t and meta["gv"]:
            cu = t.split("Giảng viên hướng dẫn:")[1].split("\n")[0].strip()
            if cu and cu != meta["gv"] and thay_mot_cum(p, cu, meta["gv"]):
                doi.append("giảng viên")
        elif t.startswith("Dự án:") and meta["de_tai"]:
            # Mẫu bìa ghi cứng tên dự án dùng chung cho ba báo cáo, nên trước
            # đây trường `Đề tài` của tệp Markdown KHÔNG lên tới bìa DOCX — đặt
            # tên đề tài xong mà bìa vẫn giữ tên cũ. Nay bìa lấy theo Markdown.
            #
            # Hai báo cáo kia khai `Đề tài` đúng bằng tên dự án đang ghi ở mẫu,
            # nên bìa của chúng không đổi một chữ nào; chỉ báo cáo nào khai tên
            # khác mới thấy khác.
            cu = t.split("Dự án:")[1].strip()
            # So sau khi gộp khoảng trắng: mẫu bìa xuống dòng giữa tên dự án,
            # còn `_lay()` đã gộp về một dòng. Không gộp thì hai chuỗi "khác
            # nhau" chỉ vì một ký tự xuống dòng, và bìa của hai báo cáo kia bị
            # viết lại — mất đúng chỗ ngắt dòng người ta đã căn.
            _gon = lambda x: re.sub(r"\s+", " ", x).strip()
            if (cu and _gon(cu) != _gon(meta["de_tai"])
                    and thay_mot_cum(p, cu, meta["de_tai"])):
                doi.append("tên đề tài")
    # Đầu trang của mẫu ghi tên học phần CỦA MẪU ("Công nghệ phần mềm") và nằm
    # ở phần header của section, không nằm trong doc.paragraphs — nên vòng lặp
    # trên không chạm tới. Hậu quả: mọi báo cáo xuất từ mẫu đều mang tên học
    # phần sai ở đầu MỌI trang, trong khi bìa thì đúng.
    if meta["hoc_phan"]:
        moi_hp = meta["hoc_phan"].replace("Học phần:", "").strip()
        da_doi_hp = False

        def _moi_doan_cua(phan):
            """Mọi đoạn trong header, KỂ CẢ đoạn nằm trong bảng.

            Mẫu đặt tên học phần trong một bảng một ô, nên `header.paragraphs`
            chỉ trả về đúng một đoạn rỗng ở ngoài bảng — quét theo nó thì không
            bao giờ chạm tới chữ cần đổi, và lỗi trôi qua im lặng.
            """
            yield from phan.paragraphs
            for bang in phan.tables:
                for dong in bang.rows:
                    for o in dong.cells:
                        yield from o.paragraphs

        for sec in doc.sections:
            for phan in (sec.header, sec.first_page_header, sec.even_page_header):
                if phan is None:
                    continue
                for p in _moi_doan_cua(phan):
                    if p.text.strip() and p.text.strip() != moi_hp:
                        thay_ca_doan(p, moi_hp)
                        da_doi_hp = True
        if da_doi_hp:
            doi.append("học phần ở đầu trang")

    # Đóng băng những thuộc tính mà trang bìa đang thừa kế từ kiểu Normal, vì
    # dung_kieu() chạy ngay sau đây sẽ sửa Normal (giãn đoạn 6pt, căn đều) và
    # làm xô lệch bố cục bìa. Giá trị đóng băng đúng bằng giá trị đang có:
    # mẫu khai <w:pPrDefault/> rỗng nên giãn đoạn hiện là 0.
    for p in doc.paragraphs:
        pf = p.paragraph_format
        if pf.space_after is None:
            pf.space_after = Pt(0)
        if pf.space_before is None:
            pf.space_before = Pt(0)
        if p.alignment is None:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Trang bìa được lấy y nguyên nên miễn trừ khỏi luật "chỉ đề mục in đậm";
    # luật đó nhắm vào phần nội dung do script dựng, không nhắm vào mẫu bìa.
    global _SO_DOAN_BIA
    _SO_DOAN_BIA = len(doc.paragraphs)
    print(f"  bìa lấy từ {MAU_BIA.name}, đã thay: {', '.join(doi) or 'không có gì'}")
    return doc


def _moi_doan(doc: Document):
    """Duyệt mọi đoạn văn trong tài liệu, kể cả đoạn nằm trong ô bảng.

    Bỏ qua các đoạn thuộc trang bìa lấy theo mẫu.
    """
    yield from doc.paragraphs[_SO_DOAN_BIA:]
    for bang in doc.tables:
        for hang in bang.rows:
            for o in hang.cells:
                yield from o.paragraphs


def gom_in_dam_ve_de_muc(doc: Document) -> int:
    """Bỏ in đậm ở mọi đoạn không phải đề mục. Trả về số run đã đổi.

    Chạy sau cùng nên phủ được cả in đậm do Markdown sinh ra lẫn in đậm do
    script tự đặt (dòng tiêu đề bảng, trang bìa).
    """
    da_doi = 0
    for p in _moi_doan(doc):
        if p.style.name.startswith("Heading"):
            continue
        if p.text.strip().upper() in DE_MUC_NGOAI_CAY:
            continue
        for r in p.runs:
            if r.bold:
                r.bold = None
                da_doi += 1
    return da_doi


def muc_luc_tu_dong(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MỤC LỤC")
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = FONT
    r.font.color.rgb = MAU_DE_MUC
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    them_truong(p, r'TOC \o "1-3" \h \z \u')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("(Mở trong Word, bấm chuột phải vào mục lục → Update Field "
                  "để cập nhật số trang.)")
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.name = FONT
    r.font.color.rgb = MUC
    doc.add_page_break()


# ----------------------------------------------------------------------------
# Bộ dựng chính
# ----------------------------------------------------------------------------

BO_QUA_TIEU_DE = ("Mục lục",)


def dung() -> Path:
    van_ban = NGUON.read_text(encoding="utf-8")
    dong = van_ban.split("\n")

    so_do = sorted(SO_DO.glob("so-do-*.png"),
                   key=lambda p: int(re.search(r"(\d+)", p.stem).group(1)))
    if not so_do:
        raise SystemExit("Chưa có sơ đồ. Chạy render_so_do.py trước.")

    # Thông tin bìa đọc từ chính khối đầu tệp Markdown, không gán cứng.
    def _lay(nhan, mac_dinh):
        """Đọc một trường ở khối bìa. `nhan` có thể là chuỗi hoặc nhiều biến thể."""
        for n in ([nhan] if isinstance(nhan, str) else nhan):
            m = re.search(rf"\*\*{n}:?\*\*\s*(.+?)(?:\n\n|\n\*\*)",
                          van_ban[:3000], flags=re.S)
            if m:
                return re.sub(r"\s+", " ", _chu_thuan(m.group(1))).strip()
        return mac_dinh

    m_loai = re.search(r"^# (BÁO CÁO .+)$", van_ban[:3000], flags=re.M)
    meta = {
        "loai": m_loai.group(1).strip() if m_loai else "BÁO CÁO",
        "hoc_phan": (m_hp.group(1).strip() if
                     (m_hp := re.search(r"^## (Học phần:.+)$", van_ban[:3000],
                                        flags=re.M)) else ""),
        "de_tai": _lay("Đề tài", "CMC Restaurant"),
        "gv": _lay(["Giảng viên hướng dẫn", "Giảng viên phụ trách",
                    "Giảng viên"], ""),
        "nhom": _lay("Nhóm thực hiện", "Nhóm 05 sinh viên"),
        "thoi_gian": _lay("Thời gian thực hiện", ""),
        "dia_diem": "Hà Nội, tháng 8 năm 2026",
    }
    dung_mau = MAU_BIA is not None and MAU_BIA.exists()
    doc = mo_theo_mau(meta) if dung_mau else Document()
    dung_kieu(doc)
    dat_trang(doc)

    if CO_TRANG_BIA and not dung_mau:
        trang_bia(doc, meta)
    muc_luc_tu_dong(doc)

    i = 0
    chi_so_so_do = 0
    cho_ve_doc = False        # đang ở trang ngang, cần quay lại khổ dọc
    trong_bia = True          # bỏ khối <div align=center> đầu file
    bo_qua_toi_tieu_de = False
    thong_ke = {"bang": 0, "hinh": 0, "ma": 0}

    while i < len(dong):
        d = dong[i]
        s = d.strip()
        ke = dong[i + 1] if i + 1 < len(dong) else ""

        # ---- bỏ khối bìa gốc của Markdown (đã dựng lại ở trang bìa) ----
        # Chỉ bỏ tới thẻ </div> đầu tiên, để giữ lại bảng thành viên ngay sau đó.
        if trong_bia:
            if s == "</div>":
                trong_bia = False
            i += 1
            continue

        # ---- bỏ khối Mục lục của Markdown (đã thay bằng field) ----
        if bo_qua_toi_tieu_de:
            if re.match(r"^#{1,3} ", s) and not s.startswith("## Mục lục"):
                bo_qua_toi_tieu_de = False
            else:
                i += 1
                continue

        if s.startswith("## Mục lục"):
            bo_qua_toi_tieu_de = True
            i += 1
            continue

        if not s or s in ("---", "<div align=\"center\">", "</div>"):
            i += 1
            continue

        # ---- tiêu đề ----
        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            cap, chu = len(m.group(1)), m.group(2).strip()
            if cap == 1:
                doc.add_page_break()
            p = doc.add_paragraph(style=f"Heading {cap}")
            viet_inline(p, chu)
            # Ép lại trên từng run cho khớp kiểu chữ đã khai báo ở dung_kieu,
            # vì viet_inline có thể đã đặt cỡ/đậm/nghiêng theo cú pháp Markdown.
            KIEU = {1: (14, True, False), 2: (13, True, False),
                    3: (13, True, False), 4: (13, True, False)}
            co, dam, nghieng = KIEU[cap]
            for r in p.runs:
                r.font.name = FONT
                r.font.size = Pt(co)
                r.bold = dam
                r.italic = nghieng
                r.font.color.rgb = MAU_DE_MUC
            i += 1
            continue

        # ---- khối mã / sơ đồ ----
        if s.startswith("```"):
            ngon_ngu = s[3:].strip()
            than = []
            i += 1
            while i < len(dong) and not dong[i].strip().startswith("```"):
                than.append(dong[i])
                i += 1
            i += 1
            if ngon_ngu in ("mermaid", "plantuml"):
                _, da_xoay = them_so_do(doc, so_do[chi_so_so_do])
                cho_ve_doc = da_xoay          # xử lý sau khi in chú thích
                chi_so_so_do += 1
                thong_ke["hinh"] += 1
            else:
                them_khoi_ma(doc, "\n".join(than))
                thong_ke["ma"] += 1
            continue

        # ---- bảng ảnh HTML ----
        if s.startswith("<table"):
            khoi = [d]
            i += 1
            while i < len(dong):
                khoi.append(dong[i])
                if "</table>" in dong[i]:
                    break
                i += 1
            i += 1
            them_thu_vien_anh(doc, "\n".join(khoi))
            thong_ke["hinh"] += 1
            continue

        # ---- ảnh <img> đơn ----
        m = re.search(r'<img[^>]*src="([^"]+)"', s)
        if m and "<table" not in s:
            them_anh(doc, (HERE / m.group(1)).resolve(),
                     rong_toi_da_cm=7.0, cao_toi_da_cm=15.0)
            thong_ke["hinh"] += 1
            i += 1
            continue

        # ---- ảnh markdown ----
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", s)
        if m:
            them_anh(doc, (HERE / m.group(2)).resolve(), cao_toi_da_cm=18.0)
            thong_ke["hinh"] += 1
            i += 1
            continue

        # ---- bảng markdown ----
        if s.startswith("|") and ke.strip().startswith("|") and re.search(r"-{3,}", ke):
            khoi = [d, ke]
            i += 2
            while i < len(dong) and dong[i].strip().startswith("|"):
                khoi.append(dong[i])
                i += 1
            them_bang(doc, khoi)
            thong_ke["bang"] += 1
            continue

        # ---- trích dẫn ----
        if s.startswith(">"):
            khoi = []
            while i < len(dong) and dong[i].strip().startswith(">"):
                khoi.append(dong[i].strip().lstrip(">").strip())
                i += 1
            them_trich_dan(doc, [x for x in khoi if x])
            continue

        # ---- danh sách ----
        m = re.match(r"^(\d+)\.\s+(.*)$", s)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing = GIAN_DONG
            p.paragraph_format.space_after = Pt(4)
            noi = [m.group(2)]
            i += 1
            while i < len(dong) and dong[i].startswith("   ") and dong[i].strip():
                noi.append(dong[i].strip())
                i += 1
            viet_inline(p, " ".join(noi))
            continue

        if re.match(r"^[-*]\s+", s):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = GIAN_DONG
            p.paragraph_format.space_after = Pt(4)
            noi = [re.sub(r"^[-*]\s+", "", s)]
            i += 1
            while i < len(dong) and dong[i].startswith("  ") and dong[i].strip() \
                    and not re.match(r"^\s*[-*]\s+", dong[i]):
                noi.append(dong[i].strip())
                i += 1
            viet_inline(p, " ".join(noi))
            continue

        # ---- chú thích hình / bảng (có thể trải nhiều dòng) ----
        if re.match(r"^\*(Hình|Bảng)\s+[\w.]+\s+—", s):
            khoi = [s]
            i += 1
            while not khoi[-1].rstrip().endswith("*") and i < len(dong):
                khoi.append(dong[i].strip())
                i += 1
            them_chu_thich(doc, " ".join(khoi).strip().strip("*"))
            if cho_ve_doc:                    # sơ đồ vừa in trên trang ngang
                doi_huong_trang(doc, ngang=False)
                cho_ve_doc = False
            continue

        # ---- đoạn văn thường ----
        khoi = [s]
        i += 1
        while i < len(dong):
            n = dong[i].strip()
            if (not n or n.startswith(("#", "|", ">", "```", "<", "!", "---"))
                    or re.match(r"^([-*]|\d+\.)\s+", n)):
                break
            khoi.append(n)
            i += 1
        p = doc.add_paragraph()
        viet_inline(p, " ".join(khoi))

    # Đặt đầu trang và số trang sau cùng, để mọi section (kể cả trang ngang)
    # đều được áp dụng. Khi lấy theo mẫu thì bỏ qua, vì đầu trang và chân trang
    # đã có sẵn trong mẫu và hàm này sẽ xóa mất ảnh trong đó.
    if not dung_mau:
        dat_dau_chan(doc, "CMC Restaurant · Báo cáo học phần INFO2005")

    if CHI_DE_MUC_DAM:
        print(f"  đã bỏ in đậm ngoài đề mục: {gom_in_dam_ve_de_muc(doc)} vị trí")

    THU_MUC_RA.mkdir(parents=True, exist_ok=True)
    dich = DICH
    try:
        doc.save(dich)
    except PermissionError:
        # Tệp đang mở trong Word. Ghi ra tên khác để không mất kết quả dựng.
        from datetime import datetime

        dich = DICH.with_name(f"{DICH.stem}_{datetime.now():%H%M%S}.docx")
        doc.save(dich)
        print(f"CẢNH BÁO: {DICH.name} đang mở trong Word nên không ghi đè được.")
        print(f"          Đã lưu sang {dich.name}. Đóng Word rồi chạy lại để ghi đúng tên.")
    print(f"Đã xuất: {dich}")
    print(f"  bảng {thong_ke['bang']} · hình {thong_ke['hinh']} "
          f"(trong đó {chi_so_so_do} sơ đồ) · khối mã {thong_ke['ma']}")
    return DICH


if __name__ == "__main__":
    dung()
